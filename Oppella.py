import streamlit as st
import pandas as pd
import os
import zipfile
import glob
import xml.etree.ElementTree as ET
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import nsdecls, qn
from docx.oxml import OxmlElement, parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


BASE_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Plantillas')

RUTAS_ENTRADA = {
    'base_maestra':     os.path.join(BASE_PATH, 'Base Maestra'),
    'nombres_clientes': os.path.join(BASE_PATH, 'Nombre Clientes'),
}

RUTAS_PLANTILLAS = {
    'excel': {
        'colombia_ecuador': os.path.join(BASE_PATH, 'Plantilla Excel', 'Plantilla CO-EC.xlsx'),
        'argentina':        os.path.join(BASE_PATH, 'Plantilla Excel', 'PLANTILLA ARGENTINA.xlsx'),
        'cac_cup':          os.path.join(BASE_PATH, 'Plantilla Excel', 'Plantilla CAC-CUP.xlsx'),
    },
    'word': {
        'espanol': os.path.join(BASE_PATH, 'Plantilla Word', 'Modelo Carta ROC.docx'),
        'ingles':  os.path.join(BASE_PATH, 'Plantilla Word', 'Modelo Carta ROC Ingles.docx'),
    }
}

CONFIGURACION_PAIS = {
    'Colombia': {'filtro_stock': 'CO', 'plantilla_excel': 'colombia_ecuador'},
    'Ecuador':  {'filtro_stock': 'EC', 'plantilla_excel': 'colombia_ecuador'},
    'Argentina':{'filtro_stock': 'AR', 'plantilla_excel': 'argentina'},
    'CAC':      {'plantilla_excel': 'cac_cup'},
    'CUP':      {'plantilla_excel': 'cac_cup'},
}

# Namespaces XML
NS = {
    'ss': 'http://schemas.openxmlformats.org/spreadsheetml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
}

for prefix, uri in NS.items():
    ET.register_namespace(prefix, uri)


# ─────────────────────────────────────────────
# HELPERS
# ─────────────────────────────────────────────

def leer_excel_desde_carpeta(carpeta, nombre_hoja=None):
    archivos = glob.glob(os.path.join(carpeta, "*.xlsx"))
    if not archivos:
        st.error(f"No se encontró ningún Excel en: {carpeta}")
        return None
    try:
        return pd.read_excel(archivos[0], sheet_name=nombre_hoja) if nombre_hoja else pd.read_excel(archivos[0])
    except Exception as e:
        st.error(f"Error al leer {archivos[0]}: {e}")
        return None


def leer_bytes_plantilla(ruta):
    if not os.path.exists(ruta):
        st.error(f"No se encuentra la plantilla: {ruta}")
        return None
    with open(ruta, 'rb') as f:
        return f.read()


# ─────────────────────────────────────────────
# MANIPULACIÓN PURA DE XML (SIN OPENPYXL)
# ─────────────────────────────────────────────

def col_letter(col_num):
    """Convierte número de columna a letra (1->A, 2->B, 27->AA)"""
    result = ""
    while col_num > 0:
        col_num -= 1
        result = chr(65 + col_num % 26) + result
        col_num //= 26
    return result


def escribir_datos_en_xlsx_puro(plantilla_bytes, df_filtro, df_agotados=None, num_hoja_base=1, pais='Colombia'):
    from lxml import etree
    import re

    NS_URI = 'http://schemas.openxmlformats.org/spreadsheetml/2006/main'
    NS_REL = 'http://schemas.openxmlformats.org/package/2006/relationships'
    output = BytesIO()

    def extraer_style_map(sheet_data, fila_ref):
        style_map = {}
        row_style = None
        for row_elem in sheet_data.findall(f'{{{NS_URI}}}row'):
            if int(row_elem.get('r', '0')) == fila_ref:
                row_style = row_elem.get('s')
                for cell in row_elem.findall(f'{{{NS_URI}}}c'):
                    ref = cell.get('r', '')
                    col_letters = ''.join(filter(str.isalpha, ref))
                    col_num = 0
                    for ch in col_letters:
                        col_num = col_num * 26 + (ord(ch) - ord('A') + 1)
                    s_val = cell.get('s')
                    if s_val:
                        style_map[col_num] = s_val
                break
        return style_map, row_style

    def crear_celda(sheet_data_row, col_idx, row_idx, value, style_map):
        cell = etree.SubElement(sheet_data_row, f'{{{NS_URI}}}c')
        cell.set('r', f'{col_letter(col_idx)}{row_idx}')
        if col_idx in style_map:
            cell.set('s', style_map[col_idx])
        if isinstance(value, (int, float)):
            v_elem = etree.SubElement(cell, f'{{{NS_URI}}}v')
            v_elem.text = str(value)
        else:
            cell.set('t', 'inlineStr')
            is_elem = etree.SubElement(cell, f'{{{NS_URI}}}is')
            t_elem  = etree.SubElement(is_elem, f'{{{NS_URI}}}t')
            t_elem.text = str(value)

    num_filas_datos    = 0
    num_columnas_datos = 0
    nombres_columnas   = []

    # ── PASO 1: Identificar qué tablas pertenecen a la hoja base ──────────────
    tablas_hoja_base = set()
    with zipfile.ZipFile(BytesIO(plantilla_bytes), 'r') as zin:
        rels_filename = f'xl/worksheets/_rels/sheet{num_hoja_base}.xml.rels'
        if rels_filename in zin.namelist():
            rels_data = zin.read(rels_filename)
            rels_root = etree.fromstring(rels_data)
            for rel in rels_root.findall(f'{{{NS_REL}}}Relationship'):
                tipo  = rel.get('Type', '')
                tgt   = rel.get('Target', '')
                if 'table' in tipo.lower() or 'table' in tgt.lower():
                    nombre = tgt.replace('../', 'xl/').replace('./', 'xl/worksheets/')
                    if not nombre.startswith('xl/'):
                        nombre = f'xl/worksheets/{nombre}'
                    nombre = nombre.replace('xl/worksheets/../', 'xl/')
                    tablas_hoja_base.add(nombre)

    # ── PASO 2: Procesar el ZIP ───────────────────────────────────────────────
    with zipfile.ZipFile(BytesIO(plantilla_bytes), 'r') as zin:
        with zipfile.ZipFile(output, 'w', zipfile.ZIP_DEFLATED) as zout:

            for item in zin.infolist():
                datos = zin.read(item.filename)

                # ── Hoja Base ─────────────────────────────────────────────────
                if item.filename == f'xl/worksheets/sheet{num_hoja_base}.xml' and not df_filtro.empty:
                    try:
                        root = etree.fromstring(datos)
                        sheet_data = root.find(f'{{{NS_URI}}}sheetData')

                        if sheet_data is not None:
                            style_map, row_style = extraer_style_map(sheet_data, fila_ref=3)

                            for row_elem in list(sheet_data.findall(f'{{{NS_URI}}}row')):
                                if int(row_elem.get('r', '0')) >= 3:
                                    sheet_data.remove(row_elem)

                            for row_idx, row_data in enumerate(df_filtro.fillna("").values.tolist(), start=3):
                                row_elem = etree.SubElement(sheet_data, f'{{{NS_URI}}}row')
                                row_elem.set('r', str(row_idx))
                                if row_style:
                                    row_elem.set('s', row_style)
                                for col_idx, value in enumerate(row_data, start=1):
                                    crear_celda(row_elem, col_idx, row_idx, value, style_map)

                            num_filas_datos    = len(df_filtro) + 2
                            num_columnas_datos = len(df_filtro.columns)
                            nombres_columnas   = df_filtro.columns.tolist()

                        datos = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
                    except Exception as e:
                        print(f"Error sheet{num_hoja_base}.xml: {e}")

                # ── Actualizar tabla SOLO si pertenece a la hoja base ──────────
                elif (item.filename.startswith('xl/tables/table') and
                      item.filename.endswith('.xml') and
                      item.filename in tablas_hoja_base and
                      num_filas_datos > 0):
                    try:
                        root = etree.fromstring(datos)

                        nuevo_rango = f"A2:{col_letter(num_columnas_datos)}{num_filas_datos}"
                        root.set('ref', nuevo_rango)
                        root.set('count', str(num_columnas_datos))

                        table_cols = root.find(f'{{{NS_URI}}}tableColumns')
                        if table_cols is not None:
                            for col_elem in list(table_cols.findall(f'{{{NS_URI}}}tableColumn')):
                                table_cols.remove(col_elem)
                            table_cols.set('count', str(num_columnas_datos))
                            for idx, col_name in enumerate(nombres_columnas, start=1):
                                col_elem = etree.SubElement(table_cols, f'{{{NS_URI}}}tableColumn')
                                col_elem.set('id', str(idx))
                                col_elem.set('name', str(col_name))

                        auto_filter = root.find(f'{{{NS_URI}}}autoFilter')
                        if auto_filter is not None:
                            auto_filter.set('ref', nuevo_rango)

                        datos = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
                    except Exception as e:
                        print(f"Error actualizando tabla {item.filename}: {e}")

                # ── Sheet7: Insertar datos de agotados (AHORA SIEMPRE) ─────────
                elif item.filename == 'xl/worksheets/sheet7.xml':
                    try:
                        root = etree.fromstring(datos)
                        sheet_data = root.find(f'{{{NS_URI}}}sheetData')

                        if sheet_data is not None:
                            style_map_ag, row_style_ag = extraer_style_map(sheet_data, fila_ref=4)

                            # Eliminar filas >= 4
                            for row_elem in list(sheet_data.findall(f'{{{NS_URI}}}row')):
                                if int(row_elem.get('r', '0')) >= 4:
                                    sheet_data.remove(row_elem)

                            # Insertar datos si df_agotados no es None y no está vacío
                            if df_agotados is not None and not df_agotados.empty:
                                for row_idx, row_data in enumerate(df_agotados.values.tolist(), start=4):
                                    row_elem = etree.SubElement(sheet_data, f'{{{NS_URI}}}row')
                                    row_elem.set('r', str(row_idx))
                                    if row_style_ag:
                                        row_elem.set('s', row_style_ag)
                                    for col_idx, value in enumerate(row_data, start=2):
                                        crear_celda(row_elem, col_idx, row_idx, value, style_map_ag)

                        datos = etree.tostring(root, xml_declaration=True, encoding='UTF-8', standalone=True)
                    except Exception as e:
                        print(f"Error sheet7.xml: {e}")

                # ── Pivots para refresh ────────────────────────────────────────
                if 'pivotTable' in item.filename and item.filename.endswith('.xml'):
                    try:
                        xml_str = datos.decode('utf-8')
                        if 'refreshOnLoad' in xml_str:
                            xml_str = re.sub(r'refreshOnLoad="[^"]*"', 'refreshOnLoad="1"', xml_str)
                        else:
                            xml_str = re.sub(r'(<pivotTableDefinition\s)', r'\1refreshOnLoad="1" ', xml_str)
                        datos = xml_str.encode('utf-8')
                    except Exception as e:
                        print(f"Error pivot {item.filename}: {e}")

                zout.writestr(item, datos)

    output.seek(0)
    return output.read()

def col_letter(col_num):
    """Convierte número de columna a letra (1->A, 2->B, 27->AA)"""
    result = ""
    while col_num > 0:
        col_num -= 1
        result = chr(65 + col_num % 26) + result
        col_num //= 26
    return result


# ─────────────────────────────────────────────
# FUNCIONES CORE
# ─────────────────────────────────────────────

def set_cell_borders(cell, border_size=1):
    cell_pr = cell._element.get_or_add_tcPr()
    tc_borders = cell_pr.find(qn('w:tcBorders'))
    if tc_borders is None:
        tc_borders = OxmlElement('w:tcBorders')
        cell_pr.append(tc_borders)
    border_attrs = {
        'top': border_size, 'left': border_size,
        'bottom': border_size, 'right': border_size,
        'insideH': border_size, 'insideV': border_size
    }
    for border, size in border_attrs.items():
        el = OxmlElement(f'w:{border}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(size * 8))
        el.set(qn('w:space'), '0')
        tc_borders.append(el)


def insert_styled_table_from_df(doc, df, cliente, fecha):
    for paragraph in doc.paragraphs:
        if '[cliente]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[cliente]', cliente)
            for run in paragraph.runs:
                run.font.size = Pt(10)
        if '[fecha]' in paragraph.text:
            paragraph.text = paragraph.text.replace('[fecha]', fecha)
            for run in paragraph.runs:
                run.font.size = Pt(10)

    for paragraph in doc.paragraphs:
        if '[tabla]' in paragraph.text:
            parts = paragraph.text.split('[tabla]')
            new_paragraph = paragraph.insert_paragraph_before()
            paragraph.text = parts[0] + parts[1]

            table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
            new_paragraph._element.addnext(table._element)
            table.autofit = False
            table.allow_autofit = False
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            column_widths = [Inches(1)] * df.shape[1]

            hdr_cells = table.rows[0].cells
            for i, col in enumerate(df.columns):
                hdr_cells[i].text = col
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                hdr_cells[i].paragraphs[0].runs[0].font.bold = True
                hdr_cells[i].paragraphs[0].runs[0].font.size = Pt(9)
                hdr_cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
                hdr_cells[i]._element.get_or_add_tcPr().append(
                    parse_xml(r'<w:shd {} w:fill="042b0b"/>'.format(nsdecls('w')))
                )
                hdr_cells[i].width = column_widths[i]

            for i in range(df.shape[0]):
                row_cells = table.rows[i + 1].cells
                for j, val in enumerate(df.iloc[i]):
                    row_cells[j].text = str(val)
                    row_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    row_cells[j].paragraphs[0].runs[0].font.size = Pt(9)
                    row_cells[j].width = column_widths[j]

            for row in table.rows:
                for cell in row.cells:
                    set_cell_borders(cell)
            return


def calcular_agotados_desde_df(df_filtro, pais, filtered_df):
    if df_filtro.empty:
        return None

    df_ag = df_filtro[
        (df_filtro['Status'] == 'Backorder') &
        (df_filtro['Sin Despacho'] > 0)
    ].copy()

    if df_ag.empty:
        return None

    df_ag = df_ag.drop_duplicates(subset='GMID').reset_index(drop=True)

    # ✅ FIX: Limpiar .0 de columnas numéricas antes de armar el resultado
    for col in ['EAN Code', 'GMID', 'PLU']:
        if col in df_ag.columns:
            df_ag[col] = pd.to_numeric(df_ag[col], errors='coerce').apply(
                lambda x: str(int(x)) if pd.notna(x) else ''
            )

    if pais in ['Colombia', 'Ecuador']:
        df_resultado = df_ag[['EAN Code', 'Canal cliente DS', 'GMID', 'PLU', 'Local Description']].rename(columns={
            'EAN Code':          'EAN',
            'Canal cliente DS':             'CANAL',
            'Local Description': 'Descripción'
        })
    elif pais == 'Argentina':
        df_resultado = df_ag[['GMID', 'Local Description']].rename(columns={
            'Local Description': 'Descripción'
        })
    else:
        df_resultado = df_ag[['Canal cliente DS', 'GMID', 'PLU', 'Local Description']].rename(columns={
            'Canal cliente DS':             'CANAL',
            'Local Description': 'Descripción'
        })

    df_resultado = df_resultado.reset_index(drop=True)
    df_resultado['Fecha Estimada de Disponibilidad'] = ''

    if filtered_df is not None and not filtered_df.empty:
        merge_f = filtered_df[['GMID', 'Available for Invoicing']].drop_duplicates().copy()

        # ✅ FIX: Normalizar GMID a string sin .0 en ambos DataFrames antes del merge
        merge_f['GMID'] = pd.to_numeric(merge_f['GMID'], errors='coerce').apply(
            lambda x: str(int(x)) if pd.notna(x) else ''
        )
        df_resultado['GMID'] = df_resultado['GMID'].astype(str).str.replace(r'\.0$', '', regex=True).str.strip()

        df_resultado = df_resultado.merge(merge_f, on='GMID', how='left')
        df_resultado['Fecha Estimada de Disponibilidad'] = (
            df_resultado['Available for Invoicing'].apply(procesar_fecha)
        )
        df_resultado.drop('Available for Invoicing', axis=1, inplace=True)

    return df_resultado if len(df_resultado) > 0 else None


def calcular_fullbox(row):
    ordered    = row.get('Ordered')
    masterpack = row.get('UNTS MASTER PACK') or row.get('Masterpack')
    if pd.isna(ordered) or pd.isna(masterpack) or masterpack == 0:
        return "Sin Datos"
    return "Full Box" if int(ordered / masterpack) == ordered / masterpack else "NO Full Box"


def procesar_fecha(val):
    if pd.isna(val):
        return None
    if val == 'Disponible':
        return 'Disponible'
    try:
        if isinstance(val, str):
            fecha = pd.to_datetime(val, format='%d/%m/%Y', errors='coerce')
            return fecha.strftime('%d/%m/%Y') if pd.notna(fecha) else None
        return val.strftime('%d/%m/%Y') if pd.notna(val) else None
    except:
        return None


def get_fecha_formateada():
    meses = {
        "January": "enero", "February": "febrero", "March": "marzo",
        "April": "abril", "May": "mayo", "June": "junio",
        "July": "julio", "August": "agosto", "September": "septiembre",
        "October": "octubre", "November": "noviembre", "December": "diciembre"
    }
    fecha = datetime.now().strftime("%d de %B de %Y")
    for en, es in meses.items():
        fecha = fecha.replace(en, es)
    return fecha


# ─────────────────────────────────────────────
# PREPARACIÓN DE DATOS
# ─────────────────────────────────────────────

def preparar_df_clientes(df_clientes_raw, pais, df_maestra_h1, df_maestra_h2, df_nombres):
    df = df_clientes_raw[df_clientes_raw['Sales Org'] == pais].reset_index(drop=True).copy()

    for col in ['Columna1','Columna2','Columna3','Columna4','Columna5',
                'Columna6','Columna7','Columna8','Columna9','Columna10',
                'Status','KeyPLU','PLU','Fullbox','Sin Despacho']:
        if col not in df.columns:
            df[col] = None

    df = pd.merge(df, df_nombres[['Sold To', 'Nombre corto']], on='Sold To', how='left')
    df['Sold To Name'] = df['Nombre corto']
    df.drop(columns=['Nombre corto'], inplace=True)

    df['Sold_To_int']  = pd.to_numeric(df['Sold To'],  errors='coerce').fillna(df['Sold To'])
    df['EAN_Code_int'] = pd.to_numeric(df['EAN Code'], errors='coerce').fillna(df['EAN Code'])

    if pais in ['Colombia', 'Ecuador']:
        df['KeyPLU'] = (
            df['Sold_To_int'].astype(int).astype(str) +
            df['EAN_Code_int'].astype(int).astype(str) +
            df['Canal cliente DS'].astype(str)
        )
    elif pais == 'Argentina':
        df['KeyPLU'] = (
            df['Sold_To_int'].fillna(0).astype(int).astype(str) +
            df['EAN_Code_int'].fillna(0).astype(int).astype(str)
        )
    else:
        try:
            df['KeyPLU'] = (
                df['Sold_To_int'].fillna(0).astype(int).astype(str) +
                df['EAN_Code_int'].fillna(0).astype(int).astype(str)
            )
        except ValueError:
            df['KeyPLU'] = None

    df = df.merge(df_maestra_h2[['Code', 'Definition']],
                  left_on='Rejection Code', right_on='Code', how='left')
    df['Status'] = df['Definition']
    df.drop(['Code', 'Definition'], axis=1, inplace=True)

    # ✅ FIX: Limpiar el KeyPLU de la maestra (puede tener .0 del Excel)
    df_maestra_h1['KeyPLU'] = (
        df_maestra_h1['KeyPLU']
        .astype(str)
        .str.replace(r'\.0', '', regex=True)
        .str.strip()
    )

    df = df.merge(df_maestra_h1[['KeyPLU', 'PLU']],
                on='KeyPLU', how='left', suffixes=('', '_maestro'))
    df['PLU'] = df['PLU_maestro'].fillna(df['PLU'])
    df.drop('PLU_maestro', axis=1, inplace=True)

    df['Fullbox']      = df.apply(calcular_fullbox, axis=1)
    df['Sin Despacho'] = df['Ordered'] - df['Invoice Sales']

    return df


def resolver_clientes_lista(modo, ids_raw, df_clientes):
    if modo == 'Todos':
        return df_clientes['Sold To Name'].dropna().unique().tolist()
    else:
        codigos = [int(c.strip()) for c in ids_raw.split(',') if c.strip().isdigit()]
        clientes_lista = []
        for codigo in codigos:
            df_filtro = df_clientes[df_clientes['Sold To'] == codigo]
            if not df_filtro.empty:
                nombre = df_filtro['Sold To Name'].iloc[0]
                if pd.notna(nombre) and nombre != '':
                    clientes_lista.append(nombre)
        return clientes_lista


# ─────────────────────────────────────────────
# PROCESAMIENTO POR CLIENTE
# ─────────────────────────────────────────────

def procesar_cliente(cliente, df_clientes, filtered_df, pais,
                     plantilla_bytes, fecha_actual, fecha_formateada):
    try:
        df_filtro = df_clientes[df_clientes['Sold To Name'] == cliente].reset_index(drop=True).copy()

        for col in df_filtro.select_dtypes(include=['datetime64']).columns:
            df_filtro[col] = df_filtro[col].dt.strftime('%Y-%m-%d %H:%M:%S')

        df_agotados = None
        if pais in ['Colombia', 'Ecuador', 'Argentina']:
            df_agotados = calcular_agotados_desde_df(df_filtro, pais, filtered_df)

        # ── Excel con manipulación pura de XML ──────────────────
        # ✅ CAMBIO: Ahora pasa también el país
        excel_bytes = escribir_datos_en_xlsx_puro(
            plantilla_bytes, 
            df_filtro, 
            df_agotados,
            num_hoja_base=1,
            pais=pais  # ← AGREGADO
        )
        archivos_dict = {f"{cliente}/{cliente}_{fecha_actual}.xlsx": excel_bytes}

        # ── Word ──────────────────────────────────────────────────
        if pais in ['Colombia', 'Ecuador'] and df_agotados is not None:
            lang       = 'ingles' if cliente == 'BRYDEN PI LIMITED TRINIDAD & TOBAGO' else 'espanol'
            word_bytes = leer_bytes_plantilla(RUTAS_PLANTILLAS['word'][lang])
            if word_bytes:
                doc = Document(BytesIO(word_bytes))
                insert_styled_table_from_df(doc, df_agotados, cliente, fecha_formateada)
                
                import tempfile, subprocess
                with tempfile.TemporaryDirectory() as tmpdir:
                    docx_path = os.path.join(tmpdir, f"{cliente}.docx")
                    pdf_path  = os.path.join(tmpdir, f"{cliente}.pdf")
                    
                    # Guardar el docx temporal
                    doc.save(docx_path)
                    
                    # Convertir a PDF con LibreOffice headless
                    result = subprocess.run(
                        ['soffice', '--headless', '--convert-to', 'pdf',
                        '--outdir', tmpdir, docx_path],
                        capture_output=True, timeout=60
                    )
                    
                    if result.returncode == 0 and os.path.exists(pdf_path):
                        with open(pdf_path, 'rb') as f:
                            archivos_dict[f"{cliente}/{cliente}_{fecha_actual}.pdf"] = f.read()
                    else:
                        # Fallback: si falla la conversión, guardar el .docx
                        buf = BytesIO()
                        doc.save(buf)
                        archivos_dict[f"{cliente}/{cliente}_{fecha_actual}.docx"] = buf.getvalue()

        return archivos_dict, None

    except Exception as e:
        import traceback
        return {}, f"{str(e)}\n{traceback.format_exc()}"


def ejecutar_procesamiento(pais, clientes_lista, df_clientes,
                           filtered_df, plantilla_bytes, progress_bar, status_text):
    fecha_actual     = datetime.now().strftime("%Y-%m-%d")
    fecha_formateada = get_fecha_formateada()
    archivos_zip     = {}
    errores          = []
    total            = len(clientes_lista)

    for idx, cliente in enumerate(clientes_lista):
        if not cliente:
            continue

        status_text.text(f"⏳ Procesando: {cliente} ({idx + 1}/{total})")
        progress_bar.progress((idx + 1) / total)

        archivos_cliente, error = procesar_cliente(
            cliente, df_clientes, filtered_df, pais,
            plantilla_bytes, fecha_actual, fecha_formateada
        )

        if error:
            errores.append(f"❌ {cliente}: {error}")
        else:
            archivos_zip.update(archivos_cliente)

    return archivos_zip, errores


# ─────────────────────────────────────────────
# UI STREAMLIT
# ─────────────────────────────────────────────

def main():
    st.set_page_config(page_title="Procesador Opella", page_icon="📦", layout="wide")
    st.title("📦 Procesador Opella — Generador de Archivos por Cliente")
    # st.caption(f"Plantillas cargadas desde: `{BASE_PATH}`")

    # with st.expander("🔍 Estado de plantillas", expanded=False):
    #     todos_ok = True
    #     checks = [
    #         ("Base Maestra",      RUTAS_ENTRADA['base_maestra']),
    #         ("Nombre Clientes",   RUTAS_ENTRADA['nombres_clientes']),
    #         ("Plantilla CO-EC",   RUTAS_PLANTILLAS['excel']['colombia_ecuador']),
    #         ("Plantilla ARG",     RUTAS_PLANTILLAS['excel']['argentina']),
    #         ("Plantilla CAC-CUP", RUTAS_PLANTILLAS['excel']['cac_cup']),
    #         ("Carta Word ES",     RUTAS_PLANTILLAS['word']['espanol']),
    #         ("Carta Word EN",     RUTAS_PLANTILLAS['word']['ingles']),
    #     ]
    #     for tipo, ruta in checks:
    #         existe = os.path.exists(ruta) if ruta.endswith(('.xlsx', '.docx')) else os.path.isdir(ruta)
    #         if not existe:
    #             todos_ok = False
    #         st.write(f"{'✅' if existe else '❌'} **{tipo}** — `{ruta}`")
    #     if not todos_ok:
    #         st.error("Algunas plantillas no se encuentran. Verifica la carpeta `Plantillas/`.")
    #     else:
    #         st.success("Todas las plantillas encontradas correctamente.")

    st.divider()

    st.subheader("1️⃣ Cargar Bases de Datos")
    col1, col2 = st.columns(2)
    with col1:
        file_stock    = st.file_uploader("📊 Base de Stock Alerta (.xlsx)", type="xlsx")
    with col2:
        file_clientes = st.file_uploader("👥 Base General de Clientes (.xlsx)", type="xlsx")

    st.subheader("2️⃣ Configurar Procesamiento")
    col3, col4 = st.columns(2)
    with col3:
        pais = st.selectbox("País", list(CONFIGURACION_PAIS.keys()))
    with col4:
        modo = st.radio("Clientes a procesar", ["Todos", "Seleccionar por ID"])

    ids_raw = ""
    if modo == "Seleccionar por ID":
        ids_raw = st.text_input(
            "IDs de clientes (separados por coma)",
            placeholder="10001, 10002, 10003"
        )

    st.divider()

    listo = file_stock is not None and file_clientes is not None
    if not listo:
        st.info("👆 Sube las dos bases para habilitar el procesamiento.")

    if st.button("🚀 Procesar", disabled=not listo, type="primary", use_container_width=True):

        with st.spinner("Leyendo archivos..."):
            df_stock_raw  = pd.read_excel(file_stock)
            df_cli_raw    = pd.read_excel(file_clientes)
            hoja2         = 'Hoja2' if pais in ['Colombia', 'Ecuador', 'Argentina'] else 'Hoja3'
            df_maestra_h1 = leer_excel_desde_carpeta(RUTAS_ENTRADA['base_maestra'], 'Hoja1')
            df_maestra_h2 = leer_excel_desde_carpeta(RUTAS_ENTRADA['base_maestra'], hoja2)
            df_nombres    = leer_excel_desde_carpeta(RUTAS_ENTRADA['nombres_clientes'])

        if any(df is None for df in [df_maestra_h1, df_maestra_h2, df_nombres]):
            st.error("No se pudieron cargar los archivos fijos. Verifica la carpeta `Plantillas/`.")
            return

        with st.spinner("Preparando datos..."):
            config_pais    = CONFIGURACION_PAIS.get(pais, {})
            filtro_stock   = config_pais.get('filtro_stock')
            plantilla_tipo = config_pais.get('plantilla_excel', 'cac_cup')

            plantilla_bytes = leer_bytes_plantilla(RUTAS_PLANTILLAS['excel'][plantilla_tipo])
            if not plantilla_bytes:
                st.error("No se pudo leer la plantilla Excel.")
                return

            filtered_df = pd.DataFrame()
            if filtro_stock:
                filtered_df = df_stock_raw[df_stock_raw['Mercado'] == filtro_stock].copy()

            df_clientes = preparar_df_clientes(
                df_cli_raw, pais, df_maestra_h1, df_maestra_h2, df_nombres
            )
            clientes_lista = resolver_clientes_lista(modo, ids_raw, df_clientes)

        if not clientes_lista:
            st.warning("No se encontraron clientes con los criterios seleccionados.")
            return

        st.info(
            f"Se procesarán **{len(clientes_lista)} clientes**: "
            f"{', '.join(str(c) for c in clientes_lista[:10])}"
            f"{'...' if len(clientes_lista) > 10 else ''}"
        )

        progress_bar = st.progress(0)
        status_text  = st.empty()

        archivos_zip, errores = ejecutar_procesamiento(
            pais            = pais,
            clientes_lista  = clientes_lista,
            df_clientes     = df_clientes,
            filtered_df     = filtered_df,
            plantilla_bytes = plantilla_bytes,
            progress_bar    = progress_bar,
            status_text     = status_text,
        )

        status_text.text("✅ ¡Procesamiento completado!")
        progress_bar.progress(1.0)

        if errores:
            with st.expander(f"⚠️ {len(errores)} advertencia(s)"):
                for e in errores:
                    st.write(e)

        if not archivos_zip:
            st.warning("No se generó ningún archivo. Verifica los datos.")
            return

        zip_buf = BytesIO()
        with zipfile.ZipFile(zip_buf, 'w', zipfile.ZIP_DEFLATED) as zf:
            for nombre_archivo, contenido in archivos_zip.items():
                zf.writestr(nombre_archivo, contenido)
        zip_buf.seek(0)

        st.success(f"📦 **{len(archivos_zip)} archivos** generados para **{pais}**.")
        st.download_button(
            label               = "📥 Descargar todos los archivos (ZIP)",
            data                = zip_buf,
            file_name           = f"Opella_{pais}_{datetime.now().strftime('%Y-%m-%d')}.zip",
            mime                = "application/zip",
            use_container_width = True,
            type                = "primary",
        )


if __name__ == "__main__":
    main()